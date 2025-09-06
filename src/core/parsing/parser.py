from docx import Document
from core.parsing.enums import ResumeSectionType, HeaderFields, ProfessionalSummaryFields, ProfessionalExperienceFields, EducationFields, ProfessionalDevelopmentFields, TechnicalSkillsFields
import re
from core.utils.helpers import TextUtils, ValidationUtils, EnumUtils, PositionUtils
from core.utils.log_helper import LogHelper
from core.services.config_manager import ConfigManager
from core.models.resume import Degree, Resume, Header, ProfessionalSummary, ProfessionalExperience, Education 
from pathlib import Path
import core.utils.paths as path_utils

class ResumeParser:
    HEADER_REQUIRED_FIELDS = [HeaderFields.NAME, HeaderFields.LOCATION, HeaderFields.PHONE, HeaderFields.EMAIL, HeaderFields.WORK_AUTHORIZED]
    PROFESSIONAL_SUMMARY_REQUIRED_FIELDS = [ProfessionalSummaryFields.SUMMARY, ProfessionalSummaryFields.HIGHLIGHTS]
    PROFESSIONAL_EXPERIENCE_POSITION_PATTERN = r"({})\s+(\d{{2}}/\d{{4}}.*?)\n"

    # TODO: add support for other file formats (pdf, txt, etc.)
    def __init__(self, resume_path: Path):
        self.config = ConfigManager()
        self.logger = LogHelper("resume_parser")
        self.logger.info(f"Initializing parser for: {resume_path}")
        self.resume_path = Path(resume_path)

        if not self.resume_path.exists():
            error_message = f"Resume file not found: {self.resume_path}"
            self.logger.error(error_message)
            raise FileNotFoundError(error_message)

        if self.resume_path.suffix.lower() != ".docx":
            error_message = f"Unsupported file format. Expected '.docx', got {self.resume_path.suffix}"
            self.logger.error(error_message)
            raise ValueError(error_message)

    def parse_resume_text(self) -> str:
        try:
            self.logger.info("Starting to parse resume")
            doc = Document(str(self.resume_path))
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            self.logger.info("Resume parsing completed successfully")
            return text
        except Exception as e:
            self.logger.error(f"Error parsing resume: {e}")
            return ""
    
    def _parse_resume_sections(self) -> dict[ResumeSectionType, list[str]]:
        try:
            self.logger.info("Starting to separate resume sections")
            doc = Document(str(self.resume_path))
            sections = {}
            current_section = ResumeSectionType.HEADER
            sections[current_section] = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    if (paragraph.style and paragraph.style.name and paragraph.style.name.startswith("Heading")) or \
                    (len(paragraph.text.strip()) < 50 and paragraph.text.strip().isupper()):
                        current_section = EnumUtils.get_section_type(paragraph.text)
                        if current_section not in sections:
                            sections[current_section] = []
                    else:
                        if current_section not in sections:
                            sections[current_section] = []
                        sections[current_section].append(paragraph.text.strip())
            self.logger.info("Separating resume sections completed successfully")
            return sections
        except Exception as e:
            self.logger.error(f"Error parsing resume sections: {e}")
            return {}

    def parse(self) -> Resume:
        sections = self._parse_resume_sections()
        return Resume(
            header = self._parse_resume_header(sections[ResumeSectionType.HEADER]),
            professional_summary = self._parse_resume_professional_summary(sections[ResumeSectionType.PROFESSIONAL_SUMMARY]),
            technical_skills = self._parse_resume_technical_skills(sections[ResumeSectionType.TECHNICAL_SKILLS]),
            professional_experience_list = self._parse_resume_professional_experience(sections[ResumeSectionType.PROFESSIONAL_EXPERIENCE]),
            education = self._parse_resume_education(sections[ResumeSectionType.EDUCATION]),
            professional_development_list = self._parse_resume_professional_development(sections[ResumeSectionType.PROFESSIONAL_DEVELOPMENT_OR_AFFILIATIONS])
        )

    def _parse_resume_header(self, header_text: list[str]) -> Header:
        if not header_text or len(header_text) < self.config.settings.parsing.min_header_lines:
            self.logger.warning(f"{ResumeSectionType.HEADER.value} section has insufficient data {header_text}")
            return {}

        header = Header()
        try :
            header.name = header_text[0]
            header.location = header_text[1]
            contact_info = header_text[2]
            if self.config.settings.parsing.header_contact_separator in contact_info:
                contact_info_parts = TextUtils.safe_split(contact_info,  self.config.settings.parsing.header_contact_separator)
                header.phone = TextUtils.safe_get_and_strip(contact_info_parts, 0)
                header.email = TextUtils.safe_get_and_strip(contact_info_parts, 1)
                header.linkedin = TextUtils.safe_get_and_strip(contact_info_parts, 2)
                header.github = TextUtils.safe_get_and_strip(contact_info_parts, 3)
            else:
                header.phone = contact_info.strip()
                header.email = None
                header.linkedin = None
                header.github = None
            header.work_authorized = header_text[3]

            self.logger.info(f"Successfully parsed {ResumeSectionType.HEADER.value} section")
            return header
        except IndexError as e:
            self.logger.error(f"Error parsing {ResumeSectionType.HEADER.value}: {e}")
            return {}

    def _parse_resume_professional_summary(self, professional_summary_text: list[str]) -> ProfessionalSummary:
        if ValidationUtils.check_if_section_empty(professional_summary_text, ResumeSectionType.PROFESSIONAL_SUMMARY):
            return {}
        
        professional_summary = ProfessionalSummary()
        try:
            professional_summary.summary = TextUtils.safe_get_and_strip(professional_summary_text, 0)
            professional_summary.highlights = professional_summary_text[1:]

            self.logger.info(f"Successfully parsed {ResumeSectionType.PROFESSIONAL_SUMMARY.value} section")
            return professional_summary
        except IndexError as e:
            self.logger.error(f"Error parsing {ResumeSectionType.PROFESSIONAL_SUMMARY.value}: {e}")
            return {}
    
    def _parse_resume_technical_skills(self, technical_skills_text: list[str]) -> list[str]:
        if ValidationUtils.check_if_section_empty(technical_skills_text, ResumeSectionType.TECHNICAL_SKILLS):
            return {}

        technical_skills: list[str] = []
        technical_skills.extend(TextUtils.clean_text_list(technical_skills_text))
        self.logger.info(f"Successfully parsed {ResumeSectionType.TECHNICAL_SKILLS.value} section")
        return technical_skills

    def _parse_resume_professional_experience(self, experience_list: list[str]) -> list[ProfessionalExperience]:
        if ValidationUtils.check_if_section_empty(experience_list, ResumeSectionType.PROFESSIONAL_EXPERIENCE):
            return []
        
        experience_text = "\n".join(experience_list)
        positions = PositionUtils.get_supported_positions(path_utils.get_positions_file_path())
        professional_experience_list: list[ProfessionalExperience] = []

        # Regex to find each position and its start line
        position_pattern =  re.compile(self.PROFESSIONAL_EXPERIENCE_POSITION_PATTERN.format('|'.join(map(re.escape, positions))))
        matches = list(position_pattern.finditer(experience_text))
        for i, match in enumerate(matches):
            professional_experience_list.append(self._parse_position_details(i, match, matches, experience_text))
        self.logger.info(f"Successfully parsed {ResumeSectionType.PROFESSIONAL_EXPERIENCE.value} section")
        return professional_experience_list

    def _parse_position_details(self, index: int, match: re.Match, matches: list[re.Match], experience_text: str) -> ProfessionalExperience:
        position = match.group(1).strip()
        dates = match.group(2).strip()

        # Capture block of experience text
        start_idx = match.end()
        end_idx = matches[index + 1].start() if index + 1 < len(matches) else len(experience_text)
        block = experience_text[start_idx:end_idx].strip()
        lines = block.splitlines()

        if len(lines) < self.config.settings.parsing.min_role_lines:
            self.logger.warning(f"Experience #{index} block has insufficient data: {block}")
            return {}

        if self.config.settings.parsing.professional_experience_company_location_separator in lines[0]:
            company_parts = TextUtils.safe_split(lines[0], self.config.settings.parsing.professional_experience_company_location_separator)
            company = TextUtils.safe_get_and_strip(company_parts, 0)
            location = TextUtils.safe_get_and_strip(company_parts, 1)
        else:
            company = TextUtils.safe_get_and_strip(lines, 0)
            location = None
        
        company_description = TextUtils.safe_get_and_strip(lines, 1)
        project_description = TextUtils.safe_get_and_strip(lines, 2) if lines[2].startswith("Project") else None
        bullet_lines = lines[2:] if not project_description else lines[3:]

        # Clean and collect bullets
        bullets = TextUtils.clean_text_list(bullet_lines)

        experience = ProfessionalExperience(
            position=position,
            dates=dates,
            company=company,
            location=location,
            company_description=company_description,
            project_description=project_description,
            bullets=bullets
        )
        return experience

    def _parse_resume_education(self, education_text: list[str]) -> Education:
        if ValidationUtils.check_if_section_empty(education_text, ResumeSectionType.EDUCATION):
            return {}

        education = Education()
        try:
            university_line = education_text[0]
            if self.config.settings.parsing.education_university_country_separator in university_line:
                parts = TextUtils.safe_split(university_line, self.config.settings.parsing.education_university_country_separator)
                education.university = TextUtils.safe_get_and_strip(parts, 0)
                education.country = TextUtils.safe_get_and_strip(parts, 1)
            else:
                education.university = university_line.strip()
                education.country = ""
                
            degree_list: list[Degree] = []
            for raw in education_text[1:]:
                degree = Degree()
                raw_parts = TextUtils.safe_split(raw, self.config.settings.parsing.education_degree_separator)
                if len(raw_parts) == 3:
                    degree.degree = TextUtils.safe_get_and_strip(raw_parts, 0)
                    degree.field_of_study = TextUtils.safe_get_and_strip(raw_parts, 1)
                    degree.year_of_graduation = TextUtils.safe_get_and_strip(raw_parts, 2)
                else:
                    degree.degree = raw.strip()
                    degree.field_of_study = ""
                    degree.year_of_graduation = ""
                degree_list.append(degree)
            education.degree_list = degree_list
            self.logger.info(f"Successfully parsed {ResumeSectionType.EDUCATION.value} section")
            return education
        except IndexError as e:
            self.logger.error(f"Error parsing {ResumeSectionType.EDUCATION.value}: {e}")
            return {}
    
    def _parse_resume_professional_development(self, professional_development_text: list[str]) -> list[str]:
        if ValidationUtils.check_if_section_empty(professional_development_text, ResumeSectionType.PROFESSIONAL_DEVELOPMENT_OR_AFFILIATIONS):
            return {}
        
        professional_development_list: list[str] = TextUtils.clean_text_list(professional_development_text)
        self.logger.info(f"Successfully parsed {ResumeSectionType.PROFESSIONAL_DEVELOPMENT_OR_AFFILIATIONS} section")
        return professional_development_list