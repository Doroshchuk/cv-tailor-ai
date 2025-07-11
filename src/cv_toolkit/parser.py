from docx import Document
from .enums import ResumeSectionType, HeaderFields, ProfessionalSummaryFields, ProfessionalExperienceFields, EducationFields, ProfessionalDevelopmentFields, TechnicalSkillsFields
import os
import re
from .utils.helpers import TextUtils, ValidationUtils, EnumUtils, PositionUtils


class ResumeParser:
    MIN_HEADER_LINES = 4
    HEADER_REQUIRED_FIELDS = [HeaderFields.NAME, HeaderFields.LOCATION, HeaderFields.PHONE, HeaderFields.EMAIL, HeaderFields.WORK_AUTHORIZED]
    MIN_ROLE_LINES = 3
    PROFESSIONAL_SUMMARY_REQUIRED_FIELDS = [ProfessionalSummaryFields.SUMMARY, ProfessionalSummaryFields.HIGHLIGHTS]
    HEADER_CONTACT_SEPARATOR = " ∙ "
    EDUCATION_DEGREE_SEPARATOR = ","
    PROFESSIONAL_EXPERIENCE_COMPANY_LOCATION_SEPARATOR = "|"
    EDUCATION_UNIVERSITY_COUNTRY_SEPARATOR = ","
    PROFESSIONAL_EXPERIENCE_POSITION_PATTERN = r"({})\s+(\d{{2}}/\d{{4}}.*?)\n"

    # TODO: add support for other file formats (pdf, txt, etc.)
    def __init__(self, resume_path: str):
        self.resume_path = resume_path

        if not os.path.exists(resume_path):
            raise FileNotFoundError(f"Resume file not found: {resume_path}")

        if not resume_path.endswith('.docx'):
            raise ValueError(f"Unsupported file format. Expected '.docx' got {resume_path}")

    def parse_resume(self) -> str:
        try:
            doc = Document(self.resume_path)
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            return text
        except Exception as e:
            print(f"Error parsing resume: {e}")
            return ""
    
    def parse_resume_sections(self) -> dict[ResumeSectionType, list[str]]:
        try:
            doc = Document(self.resume_path)
            sections = {}
            current_section = ResumeSectionType.HEADER
            sections[current_section] = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    if (paragraph.style and paragraph.style.name and paragraph.style.name.startswith("Heading")) or \
                    (len(paragraph.text.strip()) < 50 and paragraph.text.strip().isupper()):
                        current_section = EnumUtils.get_section_type(paragraph.text)
                        if current_section not in sections:
                            sections[current_section] = []
                    else:
                        if current_section not in sections:
                            sections[current_section] = []
                        sections[current_section].append(paragraph.text.strip())

            return sections
        except Exception as e:
            print(f"Error parsing resume sections: {e}")
            return {}

    def parse_resume_header(self, header_text: list[str]) -> dict[str, str]:
        if not header_text or len(header_text) < self.MIN_HEADER_LINES:
            print("Warning: Header section has insufficient data")
            return {}

        parsed_header = {}
        try :
            parsed_header[HeaderFields.NAME] = header_text[0]
            parsed_header[HeaderFields.LOCATION] = header_text[1]
            contact_info = header_text[2]
            if self.HEADER_CONTACT_SEPARATOR in contact_info:
                contact_info_parts = TextUtils.safe_split(contact_info, self.HEADER_CONTACT_SEPARATOR)
                parsed_header[HeaderFields.PHONE] = TextUtils.safe_get_and_strip(contact_info_parts, 0)
                parsed_header[HeaderFields.EMAIL] = TextUtils.safe_get_and_strip(contact_info_parts, 1)
            else:
                parsed_header[HeaderFields.PHONE] = contact_info.strip()
                parsed_header[HeaderFields.EMAIL] = None
            # parse LinkedIn, GitHub links
            parsed_header[HeaderFields.WORK_AUTHORIZED] = header_text[3]
            
            if not ValidationUtils.validate_required_fields(parsed_header, self.HEADER_REQUIRED_FIELDS, ResumeSectionType.HEADER):
                return {}

            return parsed_header
        except IndexError as e:
            print(f"Error parsing header: {e}")
            return {}

    def parse_resume_professional_summary(self, professional_summary_text: list[str]) -> dict[str, str]:
        if ValidationUtils.check_if_section_empty(professional_summary_text, ResumeSectionType.PROFESSIONAL_SUMMARY):
            return {}
        
        parsed_professional_summary = {}
        try:
            parsed_professional_summary[ProfessionalSummaryFields.SUMMARY] = TextUtils.safe_get_and_strip(professional_summary_text, 0)
            parsed_professional_summary[ProfessionalSummaryFields.HIGHLIGHTS] = professional_summary_text[1:]

            if not ValidationUtils.validate_required_fields(parsed_professional_summary, self.PROFESSIONAL_SUMMARY_REQUIRED_FIELDS, ResumeSectionType.PROFESSIONAL_SUMMARY):
                return {}

            return parsed_professional_summary
        except IndexError as e:
            print(f"Error parsing professional summary: {e}")
            return {}
    
    def parse_resume_technical_skills(self, technical_skills_text: list[str]) -> dict[str, str]:
        if ValidationUtils.check_if_section_empty(technical_skills_text, ResumeSectionType.TECHNICAL_SKILLS):
            return {}

        parsed_technical_skills = {}
        parsed_technical_skills[TechnicalSkillsFields.SKILLS] = TextUtils.clean_text_list(technical_skills_text)
        return parsed_technical_skills

    def parse_resume_professional_experience(self, experience_list: list[str]) -> list[dict]:
        if not experience_list:
            print("Warning: Education section is empty")
            return []
        
        experience_text = "\n".join(experience_list)
        positions = PositionUtils.get_supported_positions()
        parsed_experience = []

        # Regex to find each position and its start line
        position_pattern =  re.compile(self.PROFESSIONAL_EXPERIENCE_POSITION_PATTERN.format('|'.join(map(re.escape, positions))))
        matches = list(position_pattern.finditer(experience_text))
        for i, match in enumerate(matches):
            experience = self._parse_position_details(i, match, matches, experience_text)
            parsed_experience.append(experience)
        return parsed_experience

    def _parse_position_details(self, index: int, match: re.Match, matches: list[re.Match], experience_text: str) -> dict:
        position = match.group(1).strip()
        dates = match.group(2).strip()

        # Capture block of experience text
        start_idx = match.end()
        end_idx = matches[index + 1].start() if index + 1 < len(matches) else len(experience_text)
        block = experience_text[start_idx:end_idx].strip()
        lines = block.splitlines()

        if len(lines) < self.MIN_ROLE_LINES:
            print(f"Warning: Experience block has insufficient data: {block}")
            return {}

        if self.PROFESSIONAL_EXPERIENCE_COMPANY_LOCATION_SEPARATOR in lines[0]:
            company_parts = TextUtils.safe_split(lines[0], self.PROFESSIONAL_EXPERIENCE_COMPANY_LOCATION_SEPARATOR)
            company = TextUtils.safe_get_and_strip(company_parts, 0)
            location = TextUtils.safe_get_and_strip(company_parts, 1)
        else:
            company = TextUtils.safe_get_and_strip(lines, 0)
            location = None
        
        company_description = TextUtils.safe_get_and_strip(lines, 1)
        project_description = TextUtils.safe_get_and_strip(lines, 2) if lines[2].startswith("Project") else None
        bullet_lines = lines[2:] if not project_description else lines[3:]

        # Clean and collect bullets
        bullets = TextUtils.clean_text_list(bullet_lines)

        experience = {
            ProfessionalExperienceFields.POSITION: position,
            ProfessionalExperienceFields.DATES: dates,
            ProfessionalExperienceFields.COMPANY: company,
            ProfessionalExperienceFields.LOCATION: location,
            ProfessionalExperienceFields.COMPANY_DESCRIPTION: company_description,
            ProfessionalExperienceFields.PROJECT_DESCRIPTION: project_description,
            ProfessionalExperienceFields.BULLETS: bullets
        }
        return experience

    def parse_resume_education(self, education_text: list[str]) -> dict[str, str]:
        if ValidationUtils.check_if_section_empty(education_text, ResumeSectionType.EDUCATION):
            return {}

        parsed_education = {}
        try:
            university_line = education_text[0]
            if self.EDUCATION_UNIVERSITY_COUNTRY_SEPARATOR in university_line:
                parts = TextUtils.safe_split(university_line, self.EDUCATION_UNIVERSITY_COUNTRY_SEPARATOR)
                parsed_education[EducationFields.UNIVERSITY] = TextUtils.safe_get_and_strip(parts, 0)
                parsed_education[EducationFields.COUNTRY] = TextUtils.safe_get_and_strip(parts, 1)
            else:
                parsed_education[EducationFields.UNIVERSITY] = university_line.strip()
                parsed_education[EducationFields.COUNTRY] = ""
                
            degrees = []
            for raw in education_text[1:]:
                degree = {}
                raw_parts = TextUtils.safe_split(raw, self.EDUCATION_DEGREE_SEPARATOR)
                if len(raw_parts) == 3:
                    degree[EducationFields.DEGREE] = TextUtils.safe_get_and_strip(raw_parts, 0)
                    degree[EducationFields.FIELD_OF_STUDY] = TextUtils.safe_get_and_strip(raw_parts, 1)
                    degree[EducationFields.YEAR_OF_GRADUATION] = TextUtils.safe_get_and_strip(raw_parts, 2)
                else:
                    degree[EducationFields.DEGREE] = raw.strip()
                    degree[EducationFields.FIELD_OF_STUDY] = ""
                    degree[EducationFields.YEAR_OF_GRADUATION] = ""
                degrees.append(degree)
            parsed_education[EducationFields.DEGREE] = degrees
            return parsed_education
        except IndexError as e:
            print(f"Error parsing education: {e}")
            return {}
    
    def parse_resume_professional_development(self, professional_development_text: list[str]) -> dict[str, str]:
        if ValidationUtils.check_if_section_empty(professional_development_text, ResumeSectionType.PROFESSIONAL_DEVELOPMENT_OR_AFFILIATIONS):
            return {}
        
        parsed_professional_development = {}
        parsed_professional_development[ProfessionalDevelopmentFields.PROJECTS_OR_CERTIFICATIONS] = TextUtils.clean_text_list(professional_development_text)
        return parsed_professional_development

    def parse_detailed_sections(self, sections: dict[ResumeSectionType, list[str]]) -> dict:
        parsed_sections = {}
        for section_type, section_text in sections.items():
            if section_type == ResumeSectionType.HEADER:
                parsed_sections[section_type] = self.parse_resume_header(section_text)
            elif section_type == ResumeSectionType.PROFESSIONAL_SUMMARY:
                parsed_sections[section_type] = self.parse_resume_professional_summary(section_text)
            elif section_type == ResumeSectionType.TECHNICAL_SKILLS:
                parsed_sections[section_type] = self.parse_resume_technical_skills(section_text)
            elif section_type == ResumeSectionType.PROFESSIONAL_EXPERIENCE:
                parsed_sections[section_type] = self.parse_resume_professional_experience(section_text)
            elif section_type == ResumeSectionType.EDUCATION:
                parsed_sections[section_type] = self.parse_resume_education(section_text)
            elif section_type == ResumeSectionType.PROFESSIONAL_DEVELOPMENT_OR_AFFILIATIONS:
                parsed_sections[section_type] = self.parse_resume_professional_development(section_text)
            else:
                parsed_sections[section_type] = section_text
        return parsed_sections     
